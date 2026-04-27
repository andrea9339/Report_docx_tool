from __future__ import annotations

import argparse
import sys
from copy import deepcopy
from pathlib import Path
from tempfile import NamedTemporaryFile
from zipfile import ZIP_DEFLATED, ZipFile
from xml.sax.saxutils import escape

from docx import Document
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.text.paragraph import Paragraph
from openpyxl import load_workbook

TEMPLATE_NAME = "REPORT BASE LINKILLER - NOME COGNOME.docx"
NAME_PLACEHOLDER = "NOME COGNOME"
LINKS_HEADING = "B. Lista link rilevati"
NUMBERING_ID = 57
URL_FONT_NAME = "Montserrat"


def parse_args() -> argparse.Namespace:
    parser = argparse.ArgumentParser(
        description="Generate a DOCX report from an XLSX file using the Linkiller template."
    )
    parser.add_argument("xlsx_path", help="Path to the source XLSX file.")
    parser.add_argument(
        "--template",
        default=str(Path(__file__).with_name(TEMPLATE_NAME)),
        help="Optional path to the DOCX template.",
    )
    return parser.parse_args()


def read_urls_from_workbook(xlsx_path: Path) -> list[str]:
    workbook = load_workbook(xlsx_path, read_only=True, data_only=True)
    sheet = workbook.active

    try:
        header_row = next(sheet.iter_rows(min_row=1, max_row=1, values_only=True))
    except StopIteration as exc:
        raise ValueError("The XLSX file is empty.") from exc

    normalized_headers = [str(value).strip().lower() if value is not None else "" for value in header_row]
    try:
        url_index = normalized_headers.index("url")
    except ValueError as exc:
        raise ValueError("The XLSX file does not contain a 'URL' column.") from exc

    urls: list[str] = []
    for row in sheet.iter_rows(min_row=2, values_only=True):
        if url_index >= len(row):
            continue
        value = row[url_index]
        if value is None:
            continue
        text = str(value).strip()
        if text:
            urls.append(text)

    workbook.close()
    return urls


def find_paragraph_by_text(document: Document, target_text: str) -> Paragraph:
    for paragraph in document.paragraphs:
        if paragraph.text.strip() == target_text:
            return paragraph
    raise ValueError(f"Could not find paragraph '{target_text}' in the template.")


def copy_paragraph_properties(source: Paragraph, target: Paragraph) -> None:
    source_ppr = source._p.pPr
    if source_ppr is not None:
        target._p.insert(0, deepcopy(source_ppr))
    elif source.style is not None:
        target.style = source.style


def insert_paragraph_after(anchor: Paragraph, source_format: Paragraph) -> Paragraph:
    new_element = OxmlElement("w:p")
    anchor._p.addnext(new_element)
    new_paragraph = Paragraph(new_element, anchor._parent)
    copy_paragraph_properties(source_format, new_paragraph)
    return new_paragraph


def remove_paragraph(paragraph: Paragraph) -> None:
    element = paragraph._element
    parent = element.getparent()
    if parent is not None:
        parent.remove(element)


def set_numbering(paragraph: Paragraph, num_id: int) -> None:
    ppr = paragraph._p.get_or_add_pPr()
    num_pr = ppr.find(qn("w:numPr"))
    if num_pr is None:
        num_pr = OxmlElement("w:numPr")
        ppr.append(num_pr)

    ilvl = num_pr.find(qn("w:ilvl"))
    if ilvl is None:
        ilvl = OxmlElement("w:ilvl")
        num_pr.append(ilvl)
    ilvl.set(qn("w:val"), "0")

    num_id_element = num_pr.find(qn("w:numId"))
    if num_id_element is None:
        num_id_element = OxmlElement("w:numId")
        num_pr.append(num_id_element)
    num_id_element.set(qn("w:val"), str(num_id))


def clear_numbering(paragraph: Paragraph) -> None:
    ppr = paragraph._p.pPr
    if ppr is None:
        return
    num_pr = ppr.find(qn("w:numPr"))
    if num_pr is not None:
        ppr.remove(num_pr)


def set_run_font_family(run, font_name: str) -> None:
    run.font.name = font_name
    rpr = run._element.get_or_add_rPr()
    rfonts = rpr.rFonts
    if rfonts is None:
        rfonts = OxmlElement("w:rFonts")
        rpr.append(rfonts)
    for key in ("ascii", "hAnsi", "cs", "eastAsia"):
        rfonts.set(qn(f"w:{key}"), font_name)


def insert_url_list(document: Document, urls: list[str]) -> None:
    heading = find_paragraph_by_text(document, LINKS_HEADING)
    paragraphs = document.paragraphs
    heading_index = next(
        index for index, paragraph in enumerate(paragraphs) if paragraph._p is heading._p
    )
    marker = next(
        paragraph
        for paragraph in paragraphs[heading_index + 1 :]
        if paragraph.text.strip().startswith("MOTORE DI RICERCA:")
    )

    between_heading_and_marker: list[Paragraph] = []
    reached_marker = False
    for paragraph in paragraphs[heading_index + 1 :]:
        if paragraph._p is marker._p:
            reached_marker = True
            break
        between_heading_and_marker.append(paragraph)

    if not reached_marker:
        raise ValueError("Could not find the 'MOTORE DI RICERCA' paragraph in the template.")

    for paragraph in between_heading_and_marker:
        remove_paragraph(paragraph)

    before_list_spacer = insert_paragraph_after(heading, marker)
    clear_numbering(before_list_spacer)
    anchor = before_list_spacer

    for url in urls:
        list_item = insert_paragraph_after(anchor, marker)
        run = list_item.add_run(url)
        set_run_font_family(run, URL_FONT_NAME)
        set_numbering(list_item, NUMBERING_ID)
        anchor = list_item

    first_spacer = insert_paragraph_after(anchor, anchor if urls else marker)
    clear_numbering(first_spacer)

    second_spacer = insert_paragraph_after(first_spacer, first_spacer)
    clear_numbering(second_spacer)


def replace_placeholder_in_docx(docx_path: Path, placeholder: str, replacement: str) -> None:
    escaped_replacement = escape(replacement)
    with NamedTemporaryFile(delete=False, suffix=".docx", dir=docx_path.parent) as temp_file:
        temp_path = Path(temp_file.name)

    with ZipFile(docx_path, "r") as source_zip, ZipFile(temp_path, "w", compression=ZIP_DEFLATED) as target_zip:
        for entry in source_zip.infolist():
            data = source_zip.read(entry.filename)
            if entry.filename.endswith(".xml"):
                text = data.decode("utf-8")
                if placeholder in text:
                    data = text.replace(placeholder, escaped_replacement).encode("utf-8")
            target_zip.writestr(entry, data)

    temp_path.replace(docx_path)


def build_report(xlsx_path: Path, template_path: Path, output_dir: Path | None = None) -> Path:
    if not xlsx_path.is_file():
        raise FileNotFoundError(f"XLSX file not found: {xlsx_path}")
    if xlsx_path.suffix.lower() != ".xlsx":
        raise ValueError("Input file must be an .xlsx file.")
    if not template_path.is_file():
        raise FileNotFoundError(f"Template not found: {template_path}")

    subject_name = xlsx_path.stem
    destination_dir = output_dir.resolve() if output_dir is not None else xlsx_path.parent
    destination_dir.mkdir(parents=True, exist_ok=True)
    output_path = destination_dir / f"{subject_name}.docx"
    urls = read_urls_from_workbook(xlsx_path)

    document = Document(template_path)
    insert_url_list(document, urls)
    document.save(output_path)
    replace_placeholder_in_docx(output_path, NAME_PLACEHOLDER, subject_name)

    return output_path


def main() -> int:
    args = parse_args()
    xlsx_path = Path(args.xlsx_path).resolve()
    template_path = Path(args.template).resolve()

    try:
        output_path = build_report(xlsx_path, template_path)
    except Exception as exc:
        print(f"Error: {exc}", file=sys.stderr)
        return 1

    print(f"Generated: {output_path}")
    return 0


if __name__ == "__main__":
    raise SystemExit(main())
